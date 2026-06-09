"""
Ingestion & Extraction Layer
============================
Handles PDF (text + OCR fallback), DOCX, TXT resumes.
Includes robust skill extraction, cleaning, normalisation, and spaCy anonymisation.
"""

from __future__ import annotations

import io
import re
import unicodedata
import logging
from pathlib import Path
from typing import Union

logger = logging.getLogger(__name__)

# ── Optional heavy imports ────────────────────────────────────────────────────
try:
    import pdfplumber
    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False
    logger.warning("pdfplumber not installed. PDF parsing disabled.")

try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX parsing disabled.")

try:
    import pytesseract
    from PIL import Image
    _OCR_AVAILABLE = True
except ImportError:
    _OCR_AVAILABLE = False
    logger.warning("pytesseract/Pillow not installed. OCR disabled.")

try:
    import spacy
    _nlp = spacy.load("en_core_web_sm")
    _SPACY_AVAILABLE = True
except Exception:
    _nlp = None
    _SPACY_AVAILABLE = False
    logger.warning("spaCy en_core_web_sm not available. Anonymisation disabled.")


# ── Skill taxonomy ────────────────────────────────────────────────────────────
# Comprehensive multi-domain skill list for extraction
SKILL_TAXONOMY: dict[str, list[str]] = {
    "Programming Languages": [
        "python", "java", "javascript", "typescript", "c++", "c#", "c", "go", "golang",
        "rust", "kotlin", "swift", "scala", "r", "matlab", "perl", "php", "ruby",
        "bash", "shell", "powershell", "dart", "elixir", "haskell", "lua",
    ],
    "ML / AI": [
        "machine learning", "deep learning", "neural networks", "nlp",
        "natural language processing", "computer vision", "reinforcement learning",
        "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn", "xgboost",
        "lightgbm", "catboost", "huggingface", "transformers", "bert", "gpt",
        "llm", "rag", "langchain", "openai", "stable diffusion", "mlflow",
        "feature engineering", "model deployment", "model serving", "onnx",
    ],
    "Data & Analytics": [
        "pandas", "numpy", "matplotlib", "seaborn", "plotly", "tableau", "power bi",
        "looker", "dbt", "sql", "postgresql", "mysql", "sqlite", "oracle", "spark",
        "hadoop", "kafka", "airflow", "dagster", "prefect", "dask", "polars",
        "excel", "google sheets", "data warehouse", "data lake", "etl", "elt",
        "snowflake", "bigquery", "redshift", "databricks", "delta lake",
    ],
    "Web & Frontend": [
        "react", "next.js", "vue", "angular", "svelte", "html", "css", "sass",
        "tailwind", "bootstrap", "material ui", "webpack", "vite", "redux",
        "graphql", "rest api", "grpc", "websocket", "javascript", "typescript",
        "figma", "adobe xd", "responsive design", "accessibility", "wcag",
    ],
    "Backend & APIs": [
        "django", "flask", "fastapi", "spring", "spring boot", "node.js", "express",
        "nestjs", "rails", "laravel", "asp.net", "gin", "fiber", "actix",
        "microservices", "rest", "graphql", "grpc", "oauth2", "jwt",
        "rabbitmq", "redis", "celery", "kafka",
    ],
    "DevOps & Cloud": [
        "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible",
        "jenkins", "github actions", "gitlab ci", "ci/cd", "helm", "argo",
        "prometheus", "grafana", "elk", "datadog", "nginx", "linux",
        "infrastructure as code", "serverless", "lambda", "cloud native",
    ],
    "Databases": [
        "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "cassandra",
        "dynamodb", "firebase", "sqlite", "oracle", "sql server", "neo4j",
        "influxdb", "clickhouse", "supabase",
    ],
    "Soft Skills": [
        "leadership", "communication", "teamwork", "problem solving", "critical thinking",
        "agile", "scrum", "kanban", "project management", "stakeholder management",
        "mentoring", "coaching", "presentation", "negotiation",
    ],
    "Security": [
        "penetration testing", "ethical hacking", "siem", "soc", "vulnerability assessment",
        "incident response", "owasp", "iso 27001", "gdpr", "nist", "zero trust",
        "threat modeling", "firewall", "ids", "ips", "forensics", "malware analysis",
    ],
    "Finance & Business": [
        "financial modeling", "forecasting", "budgeting", "dcf", "valuation",
        "p&l", "gaap", "ifrs", "excel", "power bi", "tableau", "erp", "sap",
        "risk management", "compliance", "audit",
    ],
}

# Flat set for fast lookup (longest phrases first to avoid partial matches)
_ALL_SKILLS: list[str] = sorted(
    {skill for skills in SKILL_TAXONOMY.values() for skill in skills},
    key=len, reverse=True,
)


# ── Skill extractor ───────────────────────────────────────────────────────────
def extract_skills(text: str) -> dict[str, list[str]]:
    """
    Extract skills from resume text grouped by category.
    Returns {category: [skill, …], …} — only non-empty categories.
    """
    lower = text.lower()
    found_flat: set[str] = set()

    for skill in _ALL_SKILLS:
        # Word-boundary aware match (handle punctuation around skill names)
        pat = re.compile(
            r"(?<![a-z0-9\-])" + re.escape(skill) + r"(?![a-z0-9\-])", re.I
        )
        if pat.search(lower):
            found_flat.add(skill)

    # Group by category
    grouped: dict[str, list[str]] = {}
    for category, skills in SKILL_TAXONOMY.items():
        matched = [s for s in skills if s in found_flat]
        if matched:
            grouped[category] = matched

    return grouped


# ── Project extractor ─────────────────────────────────────────────────────────
# Section header patterns that precede project lists
_PROJECT_HEADER_RE = re.compile(
    r"(?:^|\n)\s*(?:projects?|personal projects?|academic projects?|"
    r"key projects?|notable projects?|project experience|project work|"
    r"portfolio|project highlights?)\s*[:\-–]?\s*\n",
    re.I,
)

# A project entry: starts with a bullet/dash/number OR a title-case line
_PROJECT_LINE_RE = re.compile(
    r"(?:^|\n)\s*(?:[-•▪►✓*]|\d+[.)]\s)\s*(.+?)(?=\n|$)", re.M
)

# Title-case standalone project names (2+ words, 4-80 chars, no verb endings)
_PROJECT_TITLE_RE = re.compile(
    r"(?:^|\n)((?:[A-Z][a-zA-Z0-9\-/&]+\s+){1,6}(?:[A-Z][a-zA-Z0-9\-/&]+))"
    r"(?:\s*[-–|:]|$)",
    re.M,
)

# Keyword signals that strongly suggest a project description line
_PROJECT_SIGNAL_RE = re.compile(
    r"\b(built|developed|created|implemented|designed|automated|deployed|"
    r"integrated|architected|built|engineered|optimized|led|launched|"
    r"system|app(?:lication)?|platform|tool|pipeline|model|dashboard|"
    r"website|api|service|module|framework|classifier|chatbot|bot)\b",
    re.I,
)


def extract_projects(text: str) -> list[dict]:
    """
    Extract project entries from resume text.

    Returns list of:
        {
            "title": str,       # best-guess project name
            "description": str, # raw text of the project entry
            "skills_mentioned": list[str],  # skills spotted in this blurb
        }
    """
    projects: list[dict] = []
    seen_titles: set[str] = set()

    # ── Strategy 1: find "Projects" section, extract bullet lines under it ──
    for match in _PROJECT_HEADER_RE.finditer(text):
        start = match.end()
        # Grab up to 1500 chars after header (stops at next all-caps section)
        block = text[start: start + 1500]
        next_section = re.search(r"\n[A-Z][A-Z\s]{4,}\n", block)
        if next_section:
            block = block[: next_section.start()]

        lines = [l.strip() for l in block.split("\n") if l.strip()]
        i = 0
        while i < len(lines):
            line = lines[i]
            # Skip very short or all-caps section-header-like lines
            if len(line) < 10 or (line.isupper() and len(line) < 40):
                i += 1
                continue

            # Accumulate consecutive related lines as one project blurb
            blurb_lines = [line]
            j = i + 1
            while j < len(lines) and len(lines[j]) > 5 and not re.match(
                r"^(?:education|experience|skills|certif|award|honor|language)",
                lines[j], re.I
            ):
                blurb_lines.append(lines[j])
                j += 1
                if len(blurb_lines) >= 5:
                    break

            blurb = " ".join(blurb_lines)

            if _PROJECT_SIGNAL_RE.search(blurb):
                title = _infer_title(line)
                if title.lower() not in seen_titles:
                    seen_titles.add(title.lower())
                    projects.append({
                        "title":            title,
                        "description":      blurb[:300],
                        "skills_mentioned": _skills_in_text(blurb),
                    })
            i = j if j > i + 1 else i + 1

    # ── Strategy 2: scan full text for signal lines not yet captured ──────
    if len(projects) < 2:
        for m in _PROJECT_LINE_RE.finditer(text):
            line = m.group(1).strip()
            if len(line) < 15:
                continue
            if _PROJECT_SIGNAL_RE.search(line):
                title = _infer_title(line)
                if title.lower() not in seen_titles:
                    seen_titles.add(title.lower())
                    projects.append({
                        "title":            title,
                        "description":      line[:300],
                        "skills_mentioned": _skills_in_text(line),
                    })
                    if len(projects) >= 10:
                        break

    return projects[:10]   # cap at 10


def _infer_title(line: str) -> str:
    """Pull a short title from the first part of a project line."""
    # Strip leading bullet/number
    line = re.sub(r"^[-•▪►✓*\d.)]+\s*", "", line).strip()
    # Try to get text before  ":" or "|" or "-" (common title delimiters)
    m = re.match(r"^([^:|\-–]{5,60}?)(?:\s*[:\-–|]|$)", line)
    if m:
        return m.group(1).strip()
    return line[:60].strip()


def _skills_in_text(text: str) -> list[str]:
    """Return known skills found in a short text blurb."""
    lower = text.lower()
    found = []
    for skill in _ALL_SKILLS:
        pat = re.compile(
            r"(?<![a-z0-9\-])" + re.escape(skill) + r"(?![a-z0-9\-])", re.I
        )
        if pat.search(lower):
            found.append(skill)
    return found[:10]


# ── Normalisation ─────────────────────────────────────────────────────────────
_BULLET_RE = re.compile(r"^[\u2022\u2023\u25E6\u2043\u2219•▪▸►✓✔·∙◦‣⁃]+\s*", re.M)
_MULTI_WS  = re.compile(r"[ \t]{2,}")
_MULTI_NL  = re.compile(r"\n{3,}")
_NON_ASCII = re.compile(r"[^\x00-\x7F]+")


def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFKD", text)
    text = text.replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = _BULLET_RE.sub("- ", text)
    text = _NON_ASCII.sub(" ", text)
    text = text.lower()
    text = _MULTI_WS.sub(" ", text)
    text = _MULTI_NL.sub("\n\n", text)
    return text.strip()


# ── Anonymiser ────────────────────────────────────────────────────────────────
_EMAIL_RE   = re.compile(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}", re.I)
_PHONE_RE   = re.compile(r"(\+?\d[\d\s\-().]{7,}\d)")
_URL_RE     = re.compile(r"https?://\S+|www\.\S+", re.I)
_ADDRESS_RE = re.compile(
    r"\d{1,5}\s[\w\s]{1,30}(?:street|st|avenue|ave|road|rd|lane|ln|blvd|drive|dr)\b", re.I
)


def anonymize_text(text: str) -> str:
    text = _EMAIL_RE.sub("[EMAIL]", text)
    text = _PHONE_RE.sub("[PHONE]", text)
    text = _URL_RE.sub("[URL]", text)
    text = _ADDRESS_RE.sub("[ADDRESS]", text)

    if _SPACY_AVAILABLE and _nlp:
        doc = _nlp(text[:100_000])
        replacements = []
        for ent in doc.ents:
            if ent.label_ in {"PERSON", "GPE"}:
                replacements.append((ent.start_char, ent.end_char, f"[{ent.label_}]"))
        for start, end, label in sorted(replacements, reverse=True):
            text = text[:start] + label + text[end:]

    return text


# ── PDF Parser ────────────────────────────────────────────────────────────────
def _parse_pdf(file_bytes: bytes) -> str:
    if not _PDF_AVAILABLE:
        raise RuntimeError("pdfplumber is required for PDF parsing.")

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            page_text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
            if page_text.strip():
                text_parts.append(page_text)
            elif _OCR_AVAILABLE:
                logger.info("Page %d has no text layer — running OCR …", page_num + 1)
                pil_image = page.to_image(resolution=200).original
                ocr_text  = pytesseract.image_to_string(pil_image)
                text_parts.append(ocr_text)
            else:
                logger.warning("Page %d appears image-only but OCR unavailable.", page_num + 1)

    return "\n".join(text_parts)


# ── DOCX Parser ───────────────────────────────────────────────────────────────
def _parse_docx(file_bytes: bytes) -> str:
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is required for DOCX parsing.")

    doc        = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


# ── Dispatcher ────────────────────────────────────────────────────────────────
def parse_resume(file_bytes: bytes, filename: str, anonymize: bool = True) -> dict:
    """
    Returns
    -------
    {
        filename, raw_text, clean_text, word_count,
        skills: {category: [skill,…]},
        all_skills: [skill,…],
        parse_error,
    }
    """
    ext      = Path(filename).suffix.lower()
    raw_text = ""
    error: str | None = None

    try:
        if ext == ".pdf":
            raw_text = _parse_pdf(file_bytes)
        elif ext in {".docx", ".doc"}:
            raw_text = _parse_docx(file_bytes)
        elif ext in {".txt", ".md"}:
            raw_text = file_bytes.decode("utf-8", errors="replace")
        else:
            raise ValueError(f"Unsupported file type: {ext!r}")
    except Exception as exc:
        error = str(exc)
        logger.error("Failed to parse %s: %s", filename, exc)

    clean = normalize_text(raw_text)

    # Extract skills + projects BEFORE anonymisation
    skills_grouped = extract_skills(clean)
    all_skills     = sorted({s for ss in skills_grouped.values() for s in ss})
    projects       = extract_projects(clean)

    if anonymize and clean:
        clean = anonymize_text(clean)

    return {
        "filename":    filename,
        "raw_text":    raw_text,
        "clean_text":  clean,
        "word_count":  len(clean.split()),
        "skills":      skills_grouped,
        "all_skills":  all_skills,
        "projects":    projects,
        "parse_error": error,
    }


def parse_resume_batch(
    files: list[tuple[bytes, str]],
    anonymize: bool = True,
) -> list[dict]:
    results = []
    for file_bytes, filename in files:
        parsed = parse_resume(file_bytes, filename, anonymize=anonymize)
        if parsed["parse_error"]:
            logger.error("Skipping %s due to parse error.", filename)
        results.append(parsed)
    return results